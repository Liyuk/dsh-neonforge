from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT

out = Path('docs/design/artifacts/dsh-neonforge-system-design.docx')
doc = Document(out)
body = doc._element.body
for child in list(body):
    if not child.tag.endswith('sectPr'):
        body.remove(child)

styles = doc.styles
styles['Normal'].font.name = 'Helvetica Neue'
styles['Normal'].font.size = Pt(10)
styles['Normal'].font.color.rgb = RGBColor(35, 48, 58)
for name, size, color in [('Title', 28, 'C8FF3D'), ('Heading 1', 18, '47E3D2'), ('Heading 2', 13, 'C8FF3D')]:
    s = styles[name]; s.font.name = 'Helvetica Neue'; s.font.size = Pt(size); s.font.bold = True; s.font.color.rgb = RGBColor.from_string(color)

def para(text='', style=None, bold=False, color=None):
    x = doc.add_paragraph(style=style); r = x.add_run(text); r.bold = bold
    if color: r.font.color.rgb = RGBColor.from_string(color)
    x.paragraph_format.space_after = Pt(6)
    return x

def table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers)); t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for c, v in zip(t.rows[0].cells, headers): c.text = v
    for row in rows:
        cells = t.add_row().cells
        for c, v in zip(cells, row): c.text = v; c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    return t

para('DSH NEONFORGE / SYSTEM DESIGN', 'Title')
para('后朋克 · 杂志拼贴式的赛博控制台｜从参考稿到 DSH 基础组件的可控适配', 'Subtitle', color='24303A')
para('STATUS / DESIGN-READY    SCOPE / DSH BASE UI ONLY    REV / 2026-08-31', bold=True, color='8FA3A6')
para('01 设计结论', 'Heading 1')
para('这不是典型霓虹赛博朋克，而是“后朋克 / 杂志拼贴式的赛博控制台”：深色工业底、荧光油墨、硬阴影、轻微错位、终端标签和实时状态信号。实现策略是保留 DSH 原有三栏布局、路由、交互和插件槽位，只在主题层与少量稳定语义标记上叠加 Neonforge 视觉。')
para('核心原则', 'Heading 2')
for x in ['结构不变：sidebar / conversation / details 三列和现有折叠、拖拽逻辑不改。','风格可撤销：所有规则挂在 body[data-dsh-neonforge] 下，插件卸载即恢复。','语义优先：通过 data-neonforge-* 标记定位基础组件，不依赖脆弱 DOM 层级。','可读性优先：正文使用雾灰而非纯白；酸绿只用于主动作和当前状态；代码块降低亮度。']: para(x)
para('02 视觉语言逆向拆解', 'Heading 1')
table(['维度','参考稿观察','DSH 落地规则'], [('色彩','黑 / 荧光酸绿 / 电蓝 / 青色，像荧光油墨叠印','canvas #080B10；surface #111720；action #C8FF3D；signal #47E3D2；blue #5C76FF'),('形状','卡片边缘略有切角，矩形为主，圆角克制','radius 4–10px；ticket 使用 clip-path；主按钮保留 4px 圆角'),('层次','底色、抬升面、硬阴影形成纸片叠层','raised #18212C；主选中项 5–6px 电蓝硬阴影；避免大面积 blur'),('纹理','点阵、扫描线、工业噪点作为环境，不干扰内容','frame/sidebar 使用低透明度 radial-gradient；内容区不铺纹理'),('字形','几何无衬线 + 终端标签 / 大写状态字','系统无衬线正文；状态标签使用等宽字体、字距 0.12em')])
para('03 DSH 组件映射', 'Heading 1')
for h, b in [('AppFrame / 三栏壳层','保留官方 grid-template-columns 与响应式收缩。插件在稳定容器上加 frame、sidebar、conversation、details 标记，CSS 只改背景、分隔线和局部层次。'),('WorkspaceBrowser / 工作区','工作区文件夹是信息架构锚点：普通态为深色条带；展开或选中时升为 raised card，酸绿边框 + 青色左脊 + 电蓝硬阴影。会话 ticket 采用轻微错切角，选中态使用酸绿底和深色字。'),('InputBar / 输入区','使用深色 composer card，不用白底；青色描边表达焦点，电蓝下偏移形成纸片叠层。权限、模型选择、附件和发送按钮均为方形/小圆角，主发送按钮为酸绿圆角方块。'),('Conversation / 消息','正文为 #D5E1DD，次要元信息为 #8FA3A6。代码块使用 #0D1219 + #2A3946 边框；inline code 使用 raised 色，禁止白色胶囊造成反差刺眼。'),('Details / 文件面板','右栏以深色面板和左边界分隔，保留文件树与详情卡片语义，避免把其他插件内容纳入全局重写。')]:
    para(h, 'Heading 2'); para(b)
para('04 令牌与状态', 'Heading 1')
table(['Token','值','用途','对比策略'], [('canvas','#080B10','应用底色','与正文形成 ≥ 12:1'),('surface','#111720','卡片 / 导航','与边界拉开层次'),('text','#D5E1DD','正文','避免 #FFF'),('muted','#8FA3A6','时间 / 辅助','仅用于次要信息'),('action','#C8FF3D','主动作 / 选中','深色文字置于其上'),('signal','#47E3D2','焦点 / 实时状态','描边与 tab 指示'),('blue','#5C76FF','硬阴影','不承担文本语义')])
para('05 还原为开发代码', 'Heading 1')
for x in ['主题注入：ui-theme/styles.ts 导入 neonforge.css；所有选择器以 body[data-dsh-neonforge] 为根。','语义接点：AppFrame、WorkspaceBrowser、Rows、InputBar 只增加 data 属性，不改变业务状态和事件处理。','插件装载：dsh-neonforge 作为 patch-layer bundle，通过 lib/client.js 在运行时装饰官方 DOM；link:live-dsh 仅用于本地软链开发，生产仍走 bundle。']: para(x)
para('关键验收断言：插件关闭时基础 DSH 可独立渲染；插件开启时 frame/sidebar/conversation/details 标记存在；三栏宽度、拖拽、会话选择和发送路径保持原有测试结果。')
para('06 测试驱动与验收矩阵', 'Heading 1')
table(['层级','检查项','证据'], [('单元 / 组件','AppFrame 三栏与 Neonforge 标记；InputBar 控件；Workspace ticket','Vitest focused suite'),('主题回归','CSS 仅在 data-dsh-neonforge 根下生效；代码块与白底胶囊对比度改善','styles tests + DOM assertions'),('运行时','软链插件加载、服务启动、浏览器可见 marker、点击/输入不阻塞','build plugin + browser smoke'),('代码审查','边界、可撤销性、可维护性、无跨插件全局污染','review checklist')])
para('07 设计边界与后续', 'Heading 1')
para('当前方案刻意不重做 DSH 信息架构，不引入玻璃拟态、渐变霓虹或大面积动画。后续若要增强“游戏界面”感，可在不改布局的前提下增加低频扫描线、状态角标和可选的工业纹理；必须保持正文与控件的可读性，并继续以基础组件作用域隔离。')
para('交付参考：packages/client/ui-theme/src/styles/neonforge.css · packages/client/ui-layout/src/client/AppFrame.tsx · plugins/dsh-neonforge/lib/client.js', bold=True, color='8FA3A6')
doc.save(out)
print(out)
